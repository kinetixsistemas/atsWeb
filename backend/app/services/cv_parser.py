import io
import re
from typing import List, Tuple, Optional
from pypdf import PdfReader

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

PDF_MAGIC = b'%PDF'
DOCX_MAGIC = b'PK\x03\x04'

SECTION_PATTERNS = [
    (r'^(contact|personal\s*(?:information|details)|contact\s*info)', 'contact'),
    (r'^(professional\s*summary|career\s*objective|profile|summary|about\s*me)', 'summary'),
    (r'^(work\s*experience|employment|professional\s*experience|work\s*history|career\s*history)',
     'experience'),
    (r'^(education|academic\s*background|qualifications|academic\s*qualifications)', 'education'),
    (r'^(skills|technical\s*skills|core\s*competencies|competencies|expertise)', 'skills'),
    (r'^(certifications|certificates|licenses|accreditations|professional\s*certifications)', 'certifications'),
    (r'^(languages|language\s*proficiency)', 'languages'),
    (r'^(projects|portfolio|project\s*experience)', 'projects'),
    (r'^(publications|papers|research|research\s*experience)', 'publications'),
    (r'^(references|referees)', 'references'),
]

MAX_HEADER_LENGTH = 50

MONTHS = r'(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)'
DATE_PATTERN = re.compile(
    rf'({MONTHS}\s*\d{{4}}|\d{{1,2}}[/-]\d{{4}}|\d{{4}}[/-]\d{{1,2}}|\d{{4}})\s*[-–to]+\s*({MONTHS}\s*\d{{4}}|\d{{1,2}}[/-]\d{{4}}|\d{{4}}[/-]\d{{1,2}}|\d{{4}}|present|current|now)',
    re.IGNORECASE
)
EMAIL_PATTERN = re.compile(r'[\w.+-]+@[\w-]+\.[\w.-]+')
PHONE_PATTERN = re.compile(
    r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{2,4}\)?[-.\s]?\d{2,4}[-.\s]?\d{2,4}[-.\s]?\d{0,4}'
)
LINKEDIN_PATTERN = re.compile(r'(linkedin\.com/in/[\w-]+|linkedin\.com[\w/]*)', re.IGNORECASE)
URL_PATTERN = re.compile(r'https?://(?!linkedin)[^\s]+', re.IGNORECASE)
BULLET_PATTERN = re.compile(r'^[\s]*[•\-*‣▪➢→▪▸‣⁃››]+', re.MULTILINE)
QUANTIFIABLE_PATTERN = re.compile(r'\b(\d+[%×x]|\d{3,}|increased|decreased|reduced|improved|saved|generated|led|managed|delivered|achieved|grew)\b',
                                  re.IGNORECASE)
SECTION_HEADER_PATTERN = re.compile(r'^[A-Z][A-Z\s/&]+$', re.MULTILINE)

DEGREE_KEYWORDS = [
    'bachelor', 'master', 'phd', 'doctorate', 'associate', r'b\.?s\.?', r'b\.?a\.?',
    r'm\.?s\.?', r'm\.?b\.?a\.?', r'm\.?a\.?', r'ph\.?d\.?', r'b\.?eng\.?', r'm\.?eng\.?',
    'licenciatura', 'ingeniería', 'ingenieria', 'diploma', 'degree', 'bachillerato'
]
DEGREE_PATTERN = re.compile(r'(' + '|'.join(DEGREE_KEYWORDS) + r')', re.IGNORECASE)

CERT_KEYWORDS = [
    'certified', 'certification', 'certificate', 'license', 'accredited',
    'professional', 'aws', 'azure', 'gcp', 'pmp', 'scrum', 'itil', 'cisa',
    'cissp', 'comptia', 'ccna', 'ccnp', 'ceh', 'oscp', 'cfe', 'cfa', 'cma',
    'six sigma', 'lean', 'prince2', 'toefl', 'ielts', 'dele', 'dalf'
]
CERT_PATTERN = re.compile(r'(' + '|'.join(CERT_KEYWORDS) + r')', re.IGNORECASE)

SKILL_TECH_KEYWORDS = [
    'python', 'javascript', 'typescript', 'java', 'c++', 'c#', 'ruby', 'go', 'rust',
    'swift', 'kotlin', 'php', 'scala', 'perl', 'r', 'matlab', 'sql', 'nosql',
    'react', 'angular', 'vue', 'node', 'django', 'flask', 'spring', 'express',
    'docker', 'kubernetes', 'aws', 'azure', 'gcp', 'terraform', 'ansible',
    'git', 'jenkins', 'ci/cd', 'linux', 'unix', 'rest', 'graphql', 'api',
    'html', 'css', 'sass', 'less', 'tailwind', 'bootstrap', 'redis', 'mongodb',
    'postgresql', 'mysql', 'oracle', 'mariadb', 'elasticsearch', 'kafka',
    'rabbitmq', 'nginx', 'apache', 'hadoop', 'spark', 'airflow', 'tableau',
    'power bi', 'machine learning', 'deep learning', 'nlp', 'computer vision',
    'tensorflow', 'pytorch', 'scikit-learn', 'pandas', 'numpy', 'opencv',
    'agile', 'scrum', 'jira', 'confluence', 'figma', 'sketch', 'photoshop',
    'illustrator', 'ui/ux', 'wireframe', 'prototype', 'selenium', 'cypress',
    'jest', 'mocha', 'chai', 'webpack', 'vite', 'babel', 'eslint', 'prettier',
]

LANGUAGE_NAMES = [
    'spanish', 'english', 'french', 'german', 'italian', 'portuguese', 'dutch',
    'russian', 'chinese', 'japanese', 'korean', 'arabic', 'hindi', 'bengali',
    'turkish', 'vietnamese', 'polish', 'ukrainian', 'romanian', 'czech',
    'greek', 'hungarian', 'swedish', 'norwegian', 'danish', 'finnish',
    'hebrew', 'thai', 'indonesian', 'malay', 'tagalog',
    'español', 'inglés', 'francés', 'alemán', 'italiano', 'portugués',
    'catalán', 'valenciano', 'euskera', 'gallego',
]

PROFICIENCY_LEVELS = ['native', 'fluent', 'advanced', 'intermediate', 'basic', 'bilingüe', 'bilingue', 'nativo']


# ─── File validation ───────────────────────────────────────────────

def validate_file_magic(content: bytes, filename: str) -> bool:
    ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''
    if ext == 'pdf':
        return content.startswith(PDF_MAGIC)
    elif ext in ('doc', 'docx'):
        return content.startswith(DOCX_MAGIC)
    return False


# ─── Text extraction ───────────────────────────────────────────────

def extract_text_from_pdf(content: bytes) -> str:
    reader = PdfReader(io.BytesIO(content))
    text_parts = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)
    return '\n'.join(text_parts)


def extract_text_from_docx(content: bytes) -> str:
    if DocxDocument is None:
        raise ImportError('python-docx is required for DOCX parsing')
    doc = DocxDocument(io.BytesIO(content))
    return '\n'.join(p.text for p in doc.paragraphs if p.text.strip())


def extract_text(content: bytes, filename: str) -> str:
    ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''
    if ext == 'pdf':
        return extract_text_from_pdf(content)
    elif ext in ('doc', 'docx'):
        return extract_text_from_docx(content)
    raise ValueError(f'Unsupported format: {ext}')


# ─── Section detection ─────────────────────────────────────────────

def detect_sections(text: str) -> List[Tuple[str, int, str]]:
    lines = text.split('\n')
    sections: List[Tuple[str, int, str]] = []
    section_type = 'preamble'
    section_start = 0

    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            continue

        if len(stripped) > MAX_HEADER_LENGTH:
            continue

        for pattern, stype in SECTION_PATTERNS:
            if re.match(pattern, stripped, re.IGNORECASE):
                if section_type != 'preamble':
                    sections.append((section_type, section_start, '\n'.join(lines[section_start:i])))
                section_type = stype
                section_start = i
                break

    sections.append((section_type, section_start, '\n'.join(lines[section_start:])))
    return sections


def get_section_text(sections: List[Tuple[str, int, str]], target: str) -> str:
    for stype, _, content in sections:
        if stype == target:
            return content
    return ''


# ─── Personal info extraction ──────────────────────────────────────

def extract_personal_info(text: str) -> dict:
    info = {
        'full_name': '',
        'email': '',
        'phone': '',
        'location': '',
        'linkedin': '',
        'portfolio': '',
    }

    emails = EMAIL_PATTERN.findall(text)
    if emails:
        info['email'] = emails[0]

    phones = PHONE_PATTERN.findall(text)
    if phones:
        info['phone'] = phones[0].strip()

    linkedin = LINKEDIN_PATTERN.search(text)
    if linkedin:
        info['linkedin'] = linkedin.group(0)

    urls = URL_PATTERN.findall(text)
    if urls:
        info['portfolio'] = urls[0]

    first_lines = text.split('\n')[:5]
    for line in first_lines:
        line = line.strip()
        if line and not EMAIL_PATTERN.search(line) and not PHONE_PATTERN.search(line) \
                and not LINKEDIN_PATTERN.search(line) and not re.match(r'^[\d\s\-+().]+$', line) \
                and len(line) < 80 and len(line) > 3:
            info['full_name'] = line
            break

    location_pattern = re.compile(
        r'([A-Z][a-záéíóúñ]+(?:\s+[A-Z][a-záéíóúñ]+)*(?:,\s*[A-Z]{2})?|'
        r'[A-Z][a-záéíóúñ]+\s*[-–]\s*[A-Z][a-záéíóúñ]+(?:\s+[A-Z][a-záéíóúñ]+)*)', re.MULTILINE
    )
    for line in first_lines[1:4]:
        candidates = location_pattern.findall(line)
        for c in candidates:
            if c not in (info['full_name'], info['email']) and len(c) > 5:
                info['location'] = c
                break

    return info


# ─── Summary extraction ────────────────────────────────────────────

def extract_summary(sections: List[Tuple[str, int, str]]) -> str:
    summary = get_section_text(sections, 'summary')
    if summary:
        lines = summary.split('\n')
        body = [l.strip() for l in lines[1:] if l.strip() and not re.match(r'^[A-Z\s/&]+$', l.strip())]
        return ' '.join(body) if body else ''
    return ''


# ─── Work experience extraction ────────────────────────────────────

def extract_work_experience(sections: List[Tuple[str, int, str]]) -> List[dict]:
    exp_text = get_section_text(sections, 'experience')
    if not exp_text:
        return []

    lines = exp_text.split('\n')[1:]
    entries: List[dict] = []
    current: Optional[dict] = None

    for line in lines:
        stripped = line.strip()
        if not stripped or re.match(r'^[A-Z\s/&]+$', stripped):
            continue

        date_match = DATE_PATTERN.search(stripped)
        has_bullet = bool(BULLET_PATTERN.match(stripped))

        if date_match and not has_bullet:
            if current:
                entries.append(current)
            current = {
                'company': stripped.rstrip(','),
                'position': '',
                'start_date': date_match.group(1).strip() if date_match.group(1) else '',
                'end_date': date_match.group(2).strip() if date_match.group(2) else '',
                'current': bool(re.search(r'(present|current|now)', stripped, re.IGNORECASE)),
                'description': [],
                'achievements': [],
            }
            date_range_text = date_match.group(0)
            company_part = stripped.replace(date_range_text, '').strip().rstrip(' ,-–')
            if company_part:
                current['company'] = company_part

            pos_match = re.match(
                r'^(.+?)\s+(at|@|–|-|—|,|\\|)\s+(.+?)\s*\d', stripped, re.IGNORECASE
            )
            if pos_match:
                current['position'] = pos_match.group(1).strip()
                current['company'] = pos_match.group(3).strip().rstrip(',')
        elif current:
            cleaned = BULLET_PATTERN.sub('', stripped).strip()
            is_quantifiable = bool(QUANTIFIABLE_PATTERN.search(cleaned))
            target = 'achievements' if is_quantifiable else 'description'
            current[target].append(cleaned)

    if current:
        entries.append(current)

    if not entries:
        entries = _fallback_experience_parse(lines)

    return entries


def _fallback_experience_parse(lines: List[str]) -> List[dict]:
    entries = []
    current: Optional[dict] = None
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        date_match = DATE_PATTERN.search(stripped)
        has_bullet = bool(BULLET_PATTERN.match(stripped))

        if date_match and not has_bullet:
            if current:
                entries.append(current)
            current = {
                'company': stripped.rstrip(','),
                'position': '',
                'start_date': date_match.group(1).strip(),
                'end_date': date_match.group(2).strip(),
                'current': bool(re.search(r'(present|current|now)', stripped, re.IGNORECASE)),
                'description': [],
                'achievements': [],
            }
            text_without_date = date_match.re.sub('', stripped).strip().rstrip(' ,-–')
            parts = re.split(r'\s+[|]\s+|\s+[–—-]\s+', text_without_date)
            if len(parts) >= 2:
                current['position'] = parts[0].strip()
                current['company'] = parts[1].strip()
            elif parts:
                current['company'] = parts[0]
        elif current:
            cleaned = BULLET_PATTERN.sub('', stripped).strip()
            if len(cleaned) > 5:
                is_quant = bool(QUANTIFIABLE_PATTERN.search(cleaned))
                current['achievements' if is_quant else 'description'].append(cleaned)

    if current:
        entries.append(current)
    return entries


# ─── Education extraction ──────────────────────────────────────────

def extract_education(sections: List[Tuple[str, int, str]]) -> List[dict]:
    edu_text = get_section_text(sections, 'education')
    if not edu_text:
        return []

    lines = edu_text.split('\n')[1:]
    entries: List[dict] = []
    current: Optional[dict] = None

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        has_degree = bool(DEGREE_PATTERN.search(stripped))
        date_match = DATE_PATTERN.search(stripped)
        has_bullet = bool(BULLET_PATTERN.match(stripped))

        if has_degree and not has_bullet:
            if current:
                entries.append(current)
            current = {
                'institution': stripped.rstrip(','),
                'degree': '',
                'field': '',
                'start_date': '',
                'end_date': '',
                'gpa': None,
            }
            if date_match:
                current['start_date'] = date_match.group(1).strip()
                current['end_date'] = date_match.group(2).strip() if date_match.group(2) else ''
                text_no_date = date_match.re.sub('', stripped).strip().rstrip(' ,-–')
                parts = re.split(r'\s*[|]\s*|\s*[–—-]\s*', text_no_date)
                if len(parts) >= 2:
                    current['degree'] = parts[0].strip()
                    current['institution'] = parts[1].strip()
                elif parts:
                    current['institution'] = parts[0]

            degree_match = DEGREE_PATTERN.search(stripped)
            if degree_match:
                deg = degree_match.group(1).upper()
                current['degree'] = deg

            gpa_match = re.search(r'(?:GPA|gpa|promedio)[:\s]*([\d.]+)', stripped, re.IGNORECASE)
            if gpa_match:
                current['gpa'] = gpa_match.group(1)
        elif current:
            gpa_match = re.search(r'(?:GPA|gpa|promedio)[:\s]*([\d.]+)', stripped, re.IGNORECASE)
            if gpa_match:
                current['gpa'] = gpa_match.group(1)
            elif len(stripped) > 3:
                current['field'] = stripped

    if current:
        entries.append(current)
    return entries


# ─── Skills extraction ─────────────────────────────────────────────

def extract_skills(sections: List[Tuple[str, int, str]]) -> List[str]:
    skills_text = get_section_text(sections, 'skills')
    if not skills_text:
        return []

    lines = skills_text.split('\n')[1:]
    all_skills: List[str] = []

    for line in lines:
        stripped = line.strip()
        if not stripped or re.match(r'^[A-Z\s/&]+$', stripped):
            continue
        cleaned = BULLET_PATTERN.sub('', stripped).strip()

        if ',' in cleaned or '|' in cleaned:
            items = re.split(r'[,|]', cleaned)
            for item in items:
                item = item.strip()
                if item and len(item) > 1:
                    all_skills.append(item)
        else:
            for kw in SKILL_TECH_KEYWORDS:
                if re.search(r'\b' + re.escape(kw) + r'\b', cleaned, re.IGNORECASE):
                    if kw not in all_skills:
                        all_skills.append(kw)

    seen = set()
    unique = []
    for s in all_skills:
        low = s.lower()
        if low not in seen:
            seen.add(low)
            unique.append(s)
    return unique


# ─── Certifications extraction ─────────────────────────────────────

def extract_certifications(sections: List[Tuple[str, int, str]]) -> List[dict]:
    cert_text = get_section_text(sections, 'certifications')
    if not cert_text:
        return []

    lines = cert_text.split('\n')[1:]
    entries: List[dict] = []

    for line in lines:
        stripped = line.strip()
        if not stripped or re.match(r'^[A-Z\s/&]+$', stripped):
            continue
        cleaned = BULLET_PATTERN.sub('', stripped).strip()

        entries.append({
            'name': cleaned,
            'issuer': '',
            'date': '',
            'expiration': None,
        })

        date_match = DATE_PATTERN.search(cleaned)
        if date_match:
            entries[-1]['date'] = date_match.group(1).strip()

        issuer_match = re.search(r'[-–—–]\s*(.+?)(?:\d|$)', cleaned)
        if issuer_match:
            entries[-1]['issuer'] = issuer_match.group(1).strip()

    return entries


# ─── Languages extraction ──────────────────────────────────────────

def extract_languages(sections: List[Tuple[str, int, str]]) -> List[dict]:
    lang_text = get_section_text(sections, 'languages')
    if not lang_text:
        return []

    lines = lang_text.split('\n')[1:]
    entries: List[dict] = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        cleaned = BULLET_PATTERN.sub('', stripped).strip()

        lang_found = ''
        prof_found = 'basic'

        for lang in LANGUAGE_NAMES:
            if re.search(r'\b' + re.escape(lang) + r'\b', cleaned, re.IGNORECASE):
                lang_found = lang.capitalize()
                break

        for prof in PROFICIENCY_LEVELS:
            if re.search(r'\b' + re.escape(prof) + r'\b', cleaned, re.IGNORECASE):
                prof_found = prof.lower()
                if prof_found in ('bilingüe', 'bilingue', 'nativo'):
                    prof_found = 'native'
                break

        if not lang_found and cleaned:
            parts = re.split(r'[,|–—-]', cleaned)
            lang_found = parts[0].strip().capitalize()

        if lang_found:
            entries.append({'language': lang_found, 'proficiency': prof_found})

    return entries


# ─── ATS scoring ───────────────────────────────────────────────────

def score_ats(text: str, sections: List[Tuple[str, int, str]], has_tables: bool) -> Tuple[int, List[dict]]:
    issues: List[dict] = []
    score = 0

    section_types = {s[0] for s in sections}

    for stype, label, suggestion, pts in [
        ('summary', 'Resumen Profesional', 'Incluye un resumen profesional', 10),
        ('experience', 'Experiencia Laboral', 'Agrega seccion de experiencia laboral', 10),
        ('education', 'Educacion', 'Incluye seccion de educacion', 8),
        ('skills', 'Habilidades', 'Agrega seccion de habilidades', 8),
    ]:
        if stype in section_types:
            score += pts
        else:
            issues.append({
                'section': label,
                'issue': f'Seccion "{label}" no encontrada',
                'severity': 'error',
                'recommendation': suggestion,
            })

    # Contact completeness
    contact_text = get_section_text(sections, 'contact') or text[:800]
    has_email = bool(EMAIL_PATTERN.search(contact_text))
    has_phone = bool(PHONE_PATTERN.search(contact_text))
    has_linkedin = bool(LINKEDIN_PATTERN.search(text))

    if has_email:
        score += 5
    else:
        issues.append({'section': 'Contacto', 'issue': 'Falta correo electronico', 'severity': 'error',
                       'recommendation': 'Incluye tu email profesional'})

    if has_phone:
        score += 5
    else:
        issues.append({'section': 'Contacto', 'issue': 'Falta numero de telefono', 'severity': 'warning',
                       'recommendation': 'Agrega tu telefono de contacto'})

    if has_linkedin:
        score += 5
    else:
        issues.append({'section': 'Contacto', 'issue': 'Falta perfil de LinkedIn', 'severity': 'warning',
                       'recommendation': 'Incluye la URL de tu perfil de LinkedIn'})

    # Bullet points in experience
    exp_text = get_section_text(sections, 'experience')
    if exp_text:
        bullet_lines = BULLET_PATTERN.findall(exp_text)
        if bullet_lines:
            score += 8
        else:
            issues.append({
                'section': 'Experiencia Laboral',
                'issue': 'Usa parrafos en vez de viñetas',
                'severity': 'error',
                'recommendation': 'Usa viñetas (bullet points) para describir tus logros. Los ATS procesan mejor las viñetas.',
            })

        quant = QUANTIFIABLE_PATTERN.findall(exp_text)
        if quant:
            score += 10
        else:
            issues.append({
                'section': 'Experiencia Laboral',
                'issue': 'Faltan logros cuantificables',
                'severity': 'warning',
                'recommendation': 'Agrega numeros, porcentajes y metricas a tus logros (ej: "Aumento ventas 30%").',
            })

        dates_in_exp = DATE_PATTERN.findall(exp_text)
        if dates_in_exp:
            score += 5
        else:
            issues.append({
                'section': 'Experiencia Laboral',
                'issue': 'Faltan fechas en la experiencia',
                'severity': 'error',
                'recommendation': 'Incluye fechas de inicio y fin en cada puesto.',
            })

    # Skills format
    skills_text = get_section_text(sections, 'skills')
    if skills_text:
        body = '\n'.join(skills_text.split('\n')[1:])
        has_comma_skills = bool(re.search(r'[,|]', body))
        if has_comma_skills:
            score += 5
        else:
            issues.append({
                'section': 'Habilidades',
                'issue': 'Skills no estan en formato separado por comas',
                'severity': 'info',
                'recommendation': 'Usa formato separado por comas para mejor parseo ATS.',
            })

    # Certifications
    cert_text = get_section_text(sections, 'certifications')
    if cert_text:
        score += 5
    else:
        issues.append({
            'section': 'Certificaciones',
            'issue': 'No se detectaron certificaciones',
            'severity': 'info',
            'recommendation': 'Agrega certificaciones relevantes si las tienes.',
        })

    # Languages
    lang_text = get_section_text(sections, 'languages')
    if lang_text:
        score += 5

    # File length
    num_pages = text.count('\x0c') + 1 if '\x0c' in text else max(1, len(text) // 2500)
    if num_pages <= 2:
        score += 5
    else:
        issues.append({
            'section': 'Formato',
            'issue': f'CV demasiado extenso ({num_pages} paginas)',
            'severity': 'warning',
            'recommendation': 'Manten el CV en maximo 2 paginas para mejor procesamiento ATS.',
        })

    # Tables
    if has_tables:
        score -= 5
        issues.append({
            'section': 'Formato',
            'issue': 'El CV contiene tablas',
            'severity': 'error',
            'recommendation': 'Evita usar tablas. Los ATS tienen problemas para leer informacion en tablas.',
        })

    # Section headers standard
    standard_headers = 0
    for stype, _, _ in sections:
        if stype != 'preamble' and stype != 'references':
            standard_headers += 1
    if standard_headers >= 4:
        score += 5
    else:
        issues.append({
            'section': 'Estructura',
            'issue': 'Pocas secciones estandar detectadas',
            'severity': 'warning',
            'recommendation': 'Usa secciones estandar: Experiencia, Educacion, Habilidades, etc.',
        })

    # File name check is done at endpoint level

    score = max(0, min(100, score))
    return score, issues


def has_tables_in_pdf(content: bytes) -> bool:
    try:
        reader = PdfReader(io.BytesIO(content))
        total_text = ''
        for page in reader.pages:
            total_text += page.extract_text() or ''
        lines = [l.strip() for l in total_text.split('\n') if l.strip()]
        aligned_lines = 0
        for line in lines:
            parts = line.split()
            if len(parts) >= 4 and all(len(p) > 2 for p in parts):
                aligned_lines += 1
        return aligned_lines > len(lines) * 0.6 if lines else False
    except Exception:
        return False


def has_tables_in_docx(content: bytes) -> bool:
    if DocxDocument is None:
        return False
    try:
        doc = DocxDocument(io.BytesIO(content))
        return len(doc.tables) > 0
    except Exception:
        return False


# ─── Main extraction function ──────────────────────────────────────

def extract_cv_data(content: bytes, filename: str) -> dict:
    text = extract_text(content, filename)
    sections = detect_sections(text)
    personal_info = extract_personal_info(text)
    summary = extract_summary(sections)
    work_experience = extract_work_experience(sections)
    education = extract_education(sections)
    skills = extract_skills(sections)
    certifications = extract_certifications(sections)
    languages = extract_languages(sections)

    ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''
    has_tbl = has_tables_in_pdf(content) if ext == 'pdf' else has_tables_in_docx(content)

    ats_score, structure_issues = score_ats(text, sections, has_tbl)

    return {
        'personal_info': personal_info,
        'professional_summary': summary,
        'work_experience': work_experience,
        'education': education,
        'skills': skills,
        'certifications': certifications,
        'languages': languages,
        'ats_score': ats_score,
        'structure_issues': structure_issues,
    }
